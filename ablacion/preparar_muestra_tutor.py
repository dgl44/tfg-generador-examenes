"""Prepara una muestra contextualizada de ~10 preguntas para que el TUTOR
(docente) las valore como ancla humana de la evaluación.

Genera:
  - muestra_tutor.md     -> documento limpio para enviar (contexto + rúbrica + preguntas)
  - respuestas_tutor.txt -> plantilla de respuestas
  - clave_tutor.json     -> clave oculta (etiqueta del juez) para comparar después
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
RESULTADOS = REPO / "ablacion" / "resultados_ablacion_v2.json"
JUEZ = REPO / "ablacion" / "evaluacion_juez_v2_obj.json"

import sys  # noqa: E402
sys.path.insert(0, str(REPO / "src" / "generador"))
from examen_comun import enunciado_visible  # noqa: E402

# Cuántas preguntas tomar de cada proyecto (muestra ligera, con peso en SI)
CUPOS = {
    "SI-P1 (busqueda heuristica)": 3,
    "Taxis Release 2 (SD)": 1,
    "Prog3 (POO Java)": 2,
}

CONTEXTO_PROYECTOS = """\
El código del que parten las preguntas son prácticas reales mías de la carrera:

- **Programación 3 (Java)**: prácticas de programación orientada a objetos
  (clases, herencia, interfaces, patrones de diseño).
- **Sistemas Inteligentes (Python)**: práctica de búsqueda heurística
  (implementación de A* y A*epsilon sobre un mapa).
- **Sistemas Distribuidos (Python)**: sistema de taxis que se comunican en
  tiempo real mediante sockets y Kafka (una central, taxis y clientes).
"""

RUBRICA = """\
Para cada pregunta, dime si la usarías para evaluar si el alumno comprende su
propio código, según tu criterio docente:

- **Sí**: la pondría tal cual.
- **Con retoques**: la idea es buena pero la ajustaría.
- **No**: no me parece adecuada.

No hace falta que verifiques tecnicismos ni la respuesta; me interesa tu juicio
sobre si es una buena pregunta. Una nota breve en las dudosas me viene genial.
¡Gracias!
"""

NOMBRE_LEGIBLE = {
    "SI-P1 (busqueda heuristica)": "Sistemas Inteligentes (Python)",
    "Taxis Release 2 (SD)": "Sistemas Distribuidos (Python)",
    "Prog3 (POO Java)": "Programación 3 (Java)",
}


def seleccionar(resultados, juez):
    """Selección determinista y representativa: cupo por proyecto, diversificando
    el tipo de pregunta. No se sesga por la etiqueta del juez (es una muestra para
    el juicio pedagógico del docente, no para validar al juez)."""
    seleccion = []
    for proyecto, cupo in CUPOS.items():
        items = sorted([r for r in resultados if r["proyecto"] == proyecto],
                       key=lambda r: r["nombre"])
        elegidos = []
        # 1) Diversificar tipo de pregunta
        tipos_vistos = []
        for r in items:
            if len(elegidos) >= cupo:
                break
            if r["tipo_pregunta"] not in tipos_vistos:
                elegidos.append(r)
                tipos_vistos.append(r["tipo_pregunta"])
        # 2) Completar cupo
        for r in items:
            if len(elegidos) >= cupo:
                break
            if r not in elegidos:
                elegidos.append(r)
        seleccion.extend(elegidos)
    return seleccion


def _add_codigo(doc, codigo):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for i, linea in enumerate(codigo.split("\n")):
        if i:
            run.add_break()
        run.add_text(linea)


def escribir_docx(seleccion, ruta):
    """Genera la muestra en Word, lista para que el tutor la lea y rellene."""
    doc = Document()
    doc.add_heading("Validación de preguntas generadas — TFG", level=0)
    doc.add_paragraph(
        "Hola Fidel. Te paso una muestra de preguntas que ha generado la "
        "herramienta sobre código de prácticas mías, para que me des tu criterio "
        "de docente sobre su calidad.")
    doc.add_heading("De dónde sale el código", level=1)
    for linea in CONTEXTO_PROYECTOS.replace("**", "").splitlines():
        if linea.strip():
            doc.add_paragraph(linea.strip(), style="List Bullet" if linea.strip().startswith("-") else None)
    doc.add_heading("Qué te pido", level=1)
    for linea in RUBRICA.replace("**", "").splitlines():
        if linea.strip():
            doc.add_paragraph(linea.strip())

    for i, r in enumerate(seleccion, start=1):
        doc.add_heading(f"Pregunta {i:02d}", level=2)
        doc.add_paragraph(f"Procedencia: {NOMBRE_LEGIBLE[r['proyecto']]} — {r['nombre']}")
        doc.add_paragraph("Código:")
        _add_codigo(doc, r["codigo"])
        doc.add_paragraph("Enunciado (lo que vería el alumno):")
        doc.add_paragraph(enunciado_visible(r["pregunta_generada"]))
        p = doc.add_paragraph()
        p.add_run("¿La usarías para evaluar la comprensión del alumno? (Sí / Con retoques / No): ").bold = True
        p.add_run("____________")
        doc.add_paragraph("Nota (opcional):")

    doc.save(ruta)


def main():
    resultados = [r for r in json.loads(RESULTADOS.read_text(encoding="utf-8"))
                  if "error" not in r]
    juez = {r["nombre"]: r["etiqueta_juez"]
            for r in json.loads(JUEZ.read_text(encoding="utf-8"))}

    seleccion = seleccionar(resultados, juez)

    md = ["# Validación de preguntas generadas — TFG\n",
          "Hola Fidel. Te paso una muestra de preguntas que ha generado la "
          "herramienta sobre código de prácticas mías, para que me des tu criterio "
          "de docente sobre su calidad.\n",
          "## De dónde sale el código\n", CONTEXTO_PROYECTOS, "\n## Qué te pido\n", RUBRICA,
          "\n---\n"]
    clave, plantilla = [], []
    for i, r in enumerate(seleccion, start=1):
        md.append(f"## Pregunta {i:02d}\n")
        md.append(f"*Procedencia: {NOMBRE_LEGIBLE[r['proyecto']]} — `{r['nombre']}`*\n")
        md.append("**Código:**\n")
        md.append(f"```\n{r['codigo']}\n```\n")
        md.append("**Enunciado (lo que vería el alumno):**\n")
        md.append(f"{enunciado_visible(r['pregunta_generada'])}\n")
        md.append("\n> **¿La usarías para evaluar la comprensión del alumno? (Sí / Con retoques / No):** ____________\n")
        md.append("> Nota (opcional): \n\n---\n")
        plantilla.append(f"{i:02d}=")
        clave.append({"n": i, "nombre": r["nombre"], "proyecto": r["proyecto"],
                      "lenguaje": r["lenguaje"], "tipo_pregunta": r["tipo_pregunta"],
                      "etiqueta_juez": juez.get(r["nombre"], "?")})

    escribir_docx(seleccion, REPO / "ablacion" / "muestra_tutor.docx")

    (REPO / "ablacion" / "muestra_tutor.md").write_text("\n".join(md), encoding="utf-8")
    (REPO / "ablacion" / "respuestas_tutor.txt").write_text(
        "# El tutor escribe BUENA o DEFECTUOSA tras cada =\n" + "\n".join(plantilla) + "\n",
        encoding="utf-8")
    (REPO / "ablacion" / "clave_tutor.json").write_text(
        json.dumps(clave, indent=2, ensure_ascii=False), encoding="utf-8")

    from collections import Counter
    print(f"Muestra de {len(seleccion)} preguntas para el tutor generada.")
    print("Por proyecto:", dict(Counter(r["proyecto"] for r in seleccion)))
    print("Por tipo:", dict(Counter(r["tipo_pregunta"] for r in seleccion)))
    print("Etiqueta juez (oculta):", dict(Counter(c["etiqueta_juez"] for c in clave)))


if __name__ == "__main__":
    main()
