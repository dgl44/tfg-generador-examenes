"""Ronda 2 de validación para el tutor: muestra afinada de Sistemas Inteligentes.

Toma SOLO el código propio del alumno (descartando el de la plantilla mediante
comparación de firmas), genera preguntas con el prompt ya afinado y produce un
Word listo para enviar. Uso: python ablacion/ronda2_tutor.py
"""

import sys
from pathlib import Path

import truststore  # noqa: E402
truststore.inject_into_ssl()

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO / "src" / "generador"))

from dotenv import load_dotenv  # noqa: E402
load_dotenv(REPO / ".env")

from docx import Document  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

from parser import extraer_unidades, firma_unidad  # noqa: E402
from modelos import ContextoAcademico, TipoPregunta  # noqa: E402
from agentes import construir_crew_para_unidad  # noqa: E402
from pipeline import _construir_resultado  # noqa: E402
from examen_comun import enunciado_visible  # noqa: E402

PLANTILLA = Path(r"C:/Users/Damian/Documents/UA_INFORMATICA/3º/SI/P1plantilla")
ALUMNO = Path(r"C:/Users/Damian/Documents/UA_INFORMATICA/3º/SI/SI-P1-ENTREGAFINAL/Fuente")
SALIDA = REPO / "ablacion" / "muestra_tutor_ronda2.docx"
MAX_LINEAS = 80

CONTEXTO = ContextoAcademico(
    asignatura="Sistemas Inteligentes",
    curso="Grado",
    titulacion="Ingeniería Informática",
    nivel="intermedio",
    usar_revisor=True,
)
TIPOS = [TipoPregunta.TEST, TipoPregunta.TRAZA, TipoPregunta.ABIERTA]

RUBRICA = [
    "Para cada pregunta, dime si la usarías para evaluar si el alumno comprende su "
    "propio código, según tu criterio docente:",
    "- Sí: la pondría tal cual.",
    "- Con retoques: la idea es buena pero la ajustaría.",
    "- No: no me parece adecuada.",
    "No hace falta que verifiques tecnicismos ni la respuesta; me interesa tu juicio "
    "sobre si es una buena pregunta. Una nota breve en las dudosas me viene genial. ¡Gracias!",
]


def unidades_propias():
    firmas_plantilla = set()
    for f in PLANTILLA.rglob("*.py"):
        for u in extraer_unidades(f):
            firmas_plantilla.add(firma_unidad(u))

    vistas, propias = set(), []
    for f in sorted(ALUMNO.rglob("*.py")):
        for u in extraer_unidades(f):
            if u.nombre in vistas:
                continue
            vistas.add(u.nombre)
            lineas = u.linea_fin - u.linea_inicio + 1
            if u.es_trivial() or lineas > MAX_LINEAS:
                continue
            if firma_unidad(u) in firmas_plantilla:
                continue  # código de plantilla: se descarta
            propias.append(u)
    return propias


def _add_codigo(doc, codigo):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for i, linea in enumerate(codigo.split("\n")):
        if i:
            run.add_break()
        run.add_text(linea)


def main():
    unidades = unidades_propias()
    print(f"Unidades propias (sin plantilla, no triviales): {len(unidades)}")

    doc = Document()
    doc.add_heading("Validación de preguntas generadas (2ª muestra) — TFG", level=0)
    doc.add_paragraph(
        "Hola Fidel. Tras tu feedback anterior, la herramienta ahora descarta el "
        "código de la plantilla y genera preguntas solo sobre lo que ha escrito el "
        "alumno. Esta muestra es solo de Sistemas Inteligentes y solo con mi propio "
        "código de la práctica de búsqueda heurística (A* y A*epsilon).")
    doc.add_heading("Qué te pido", level=1)
    for linea in RUBRICA:
        doc.add_paragraph(linea)

    for i, unidad in enumerate(unidades, start=1):
        tipo = TIPOS[(i - 1) % len(TIPOS)]
        print(f"[{i}/{len(unidades)}] {unidad.nombre} [{tipo.value}]")
        crew = construir_crew_para_unidad(unidad, CONTEXTO, tipo)
        salida = crew.kickoff()
        r = _construir_resultado(unidad, tipo, salida)
        enunciado = enunciado_visible(r.pregunta_generada)

        doc.add_heading(f"Pregunta {i:02d}", level=2)
        doc.add_paragraph(f"Procedencia: Sistemas Inteligentes (Python) — {unidad.nombre}")
        doc.add_paragraph("Código:")
        _add_codigo(doc, unidad.codigo)
        doc.add_paragraph("Enunciado (lo que vería el alumno):")
        doc.add_paragraph(enunciado)
        p = doc.add_paragraph()
        p.add_run("¿La usarías para evaluar la comprensión del alumno? (Sí / Con retoques / No): ").bold = True
        p.add_run("____________")
        doc.add_paragraph("Nota (opcional):")

    doc.save(SALIDA)
    print(f"\n>>> Muestra guardada en {SALIDA}")


if __name__ == "__main__":
    main()
