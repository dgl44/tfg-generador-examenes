"""Generación del examen en Word (.docx) a partir de las preguntas (RF10).

Produce un documento editable para que el docente lo adapte a su gusto. El
contenido es el del examen del alumno (enunciado + código), igual que el PDF;
la solución se recorta (ver :mod:`examen_comun`).
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from examen_comun import enunciado_visible, incluye_codigo, titulo_examen
from modelos import ContextoAcademico, ResultadoPregunta

_GRIS_CODIGO = RGBColor(0x33, 0x33, 0x33)


def _añadir_codigo(doc: Document, codigo: str) -> None:
    """Añade el código en una fuente monoespaciada conservando los saltos."""
    parrafo = doc.add_paragraph()
    parrafo.paragraph_format.space_after = Pt(6)
    run = parrafo.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = _GRIS_CODIGO
    for i, linea in enumerate(codigo.split("\n")):
        if i:
            run.add_break()
        run.add_text(linea)


def generar_examen_docx(
    resultados: list[ResultadoPregunta],
    contexto: ContextoAcademico | None = None,
    titulo: str | None = None,
) -> bytes:
    """Genera el examen en Word con las preguntas dadas y lo devuelve en bytes."""
    asignatura = contexto.asignatura if contexto else "Examen"
    doc = Document()

    doc.add_heading(titulo_examen(titulo, asignatura), level=0)
    doc.add_paragraph("Nombre y apellidos: ____________________________________")
    doc.add_paragraph("Fecha: __________________")

    for i, r in enumerate(resultados, start=1):
        enunciado = enunciado_visible(r.pregunta_generada)
        doc.add_heading(f"Pregunta {i} ({r.tipo.value})", level=2)
        doc.add_paragraph(enunciado)
        # Código de la unidad, salvo que el enunciado ya lo incluya
        if not incluye_codigo(enunciado, r.unidad.codigo):
            doc.add_paragraph(f"Código ({r.unidad.nombre}):")
            _añadir_codigo(doc, r.unidad.codigo)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
